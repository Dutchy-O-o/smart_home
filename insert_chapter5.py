"""
Insert the new Chapter 5 implementation sub-sections (text, figure
placeholders, embedded Mermaid PNGs) into the existing project report .docx,
and ALSO patch the existing chapter:

  - Append two new rows to Table 5.2 (POST/DELETE automations endpoints)
  - Insert two new endpoint description blocks (saveAutomation,
    deleteAutomation) right before the "Edge Layer & MQTT Topic
    Architecture" sub-section, so they sit alongside the other Device /
    Telemetry / Automation endpoints.

The new chapter sections include detailed request/response payload bodies
for every endpoint the mobile app actually invokes — both AWS and external
services (Anthropic, Spotify, Raspberry Pi).

Style conventions inferred from the existing document:
  - Heading 2: top-level sub-section
  - Heading 3: nested sub-section
  - Heading 4: leaf sub-section / endpoint
  - Normal:  body text, code/JSON blocks
  - Caption: figure / table titles
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm


SRC = "2022510085_2022510105_2022510111_project_report_4.docx"
OUT = "2022510085_2022510105_2022510111_project_report_10.docx"


# ---------------------------------------------------------------------------
# Formatting helpers (match the existing chapter conventions)
#
# Body paragraph format (taken from the report's existing Chapter 5):
#   Alignment:        Justified
#   First-line indent: 0.5 cm
#   Spacing:          Before 12 pt, After 24 pt
#   Line spacing:     1.5
#
# Code / JSON blocks (overrides):
#   Alignment:        Left
#   No first-line indent
#   Tight spacing:    Before 0 pt, After 0 pt, Single line spacing
#   Font:             Consolas 10 pt
#
# Heading numbering:
#   Existing chapter uses a multilevel list (numId=11). Our new headings must
#   reuse the same list at the correct ilvl so Word's auto-numbering stays
#   continuous: Heading 2 -> ilvl 1, Heading 3 -> ilvl 2, Heading 4 -> ilvl 3.
# ---------------------------------------------------------------------------

NUM_ID = "11"
HEADING_ILVL = {"Heading 2": "1", "Heading 3": "2", "Heading 4": "3"}


def apply_body_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0.5)
    pf.space_before = Pt(12)
    pf.space_after = Pt(24)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def apply_code_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for run in paragraph.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        # Ensure East-Asian font slot also gets Consolas so Word doesn't fall back
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "cs"):
            rFonts.set(qn(f"w:{attr}"), "Consolas")


def apply_caption_format(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def apply_heading_numbering(paragraph, style_name):
    """Attach the multilevel-list pointer (numId=11) so the existing
    Word auto-numbering picks up the new heading."""
    ilvl_val = HEADING_ILVL.get(style_name)
    if ilvl_val is None:
        return  # Heading 1 / unknown — leave alone
    pPr = paragraph._element.get_or_add_pPr()
    # Remove any existing numPr first so re-runs stay idempotent
    for existing in pPr.findall(qn("w:numPr")):
        pPr.remove(existing)
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), ilvl_val)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), NUM_ID)
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def find_anchor(doc):
    for p in doc.paragraphs:
        if p.style.name == "Heading 1" and p.text.strip() == "REFERENCES":
            return p._element
    raise RuntimeError("REFERENCES heading not found")


def find_paragraph_by_text(doc, style_name, contains):
    for p in doc.paragraphs:
        if p.style.name == style_name and contains in p.text:
            return p._element
    raise RuntimeError(f"Paragraph not found: style={style_name} contains={contains!r}")


def find_table_5_2(doc):
    """Table 5.2 has the header 'Method | Endpoint | Lambda | Summary'."""
    for tbl in doc.tables:
        cells = tbl.rows[0].cells
        if (len(cells) == 4
                and cells[0].text.strip() == "Method"
                and "Endpoint" in cells[1].text
                and "Lambda" in cells[2].text
                and "Summary" in cells[3].text):
            return tbl
    raise RuntimeError("Table 5.2 not found")


def append_row(table, values):
    row = table.add_row()
    for i, v in enumerate(values):
        row.cells[i].text = str(v)


def insert_paragraph_at(doc, anchor, text, style="Normal", role="body"):
    """
    role:
      'body'    - body text: justify, first-line indent, 1.5 spacing
      'code'    - code/JSON line: left-aligned, monospace, tight spacing
      'caption' - figure / table caption: centered, single spacing
      'raw'     - leave whatever Word's style defines, no override
    """
    p = doc.add_paragraph(text, style=style)
    el = p._element
    el.getparent().remove(el)
    anchor.addprevious(el)

    if style.startswith("Heading"):
        apply_heading_numbering(p, style)
    elif role == "body" and style == "Normal":
        apply_body_format(p)
    elif role == "code":
        apply_code_format(p)
    elif role == "caption" or style == "Caption":
        apply_caption_format(p)
    return p


def insert_table(doc, anchor, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = str(val)
    el = table._element
    el.getparent().remove(el)
    anchor.addprevious(el)
    return table


def _set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_borders_dashed(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "999999")
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def insert_image(doc, anchor, image_path, fig_number, caption, width_inches=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    el = p._element
    el.getparent().remove(el)
    anchor.addprevious(el)

    cap = doc.add_paragraph(f"Figure {fig_number} – {caption}", style="Caption")
    apply_caption_format(cap)
    cap_el = cap._element
    cap_el.getparent().remove(cap_el)
    anchor.addprevious(cap_el)


def insert_figure_placeholder(doc, anchor, fig_number, description,
                              caption, kind="screenshot"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    _set_cell_shading(cell, "FFF7E0")
    _set_cell_borders_dashed(cell)
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = 1
    run1 = p1.add_run(f"[ INSERT FIGURE {fig_number} HERE ]")
    run1.bold = True
    run1.font.size = Pt(11)
    run1.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

    p2 = cell.add_paragraph()
    p2.alignment = 1
    run2 = p2.add_run(description)
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    cell.add_paragraph()

    el = table._element
    el.getparent().remove(el)
    anchor.addprevious(el)

    cap = doc.add_paragraph(f"Figure {fig_number} – {caption}", style="Caption")
    apply_caption_format(cap)
    cap_el = cap._element
    cap_el.getparent().remove(cap_el)
    anchor.addprevious(cap_el)


# ---------------------------------------------------------------------------
# Existing-chapter patches: Table 5.2 rows + 2 new endpoint blocks
# ---------------------------------------------------------------------------

def patch_existing_chapter(doc):
    # 1) Add two rows to Table 5.2
    t52 = find_table_5_2(doc)
    append_row(t52, [
        "POST",
        "/{home id}/automations",
        "saveAutomation",
        "Persists a new sensor- or emotion-triggered automation rule.",
    ])
    append_row(t52, [
        "DELETE",
        "/{home id}/automations?rule_id=...",
        "deleteAutomation",
        "Removes the automation rule with the supplied UUID.",
    ])

    # 2) Insert new endpoint blocks BEFORE "Edge Layer & MQTT Topic Architecture"
    edge_anchor = find_paragraph_by_text(
        doc, "Heading 2", "Edge Layer & MQTT Topic Architecture"
    )

    # New H4 + body for saveAutomation
    insert_paragraph_at(
        doc, edge_anchor,
        "POST /{home_id}/automations — saveAutomation",
        style="Heading 4",
    )
    insert_paragraph_at(
        doc, edge_anchor,
        "This endpoint persists a new automation rule scoped to the specified "
        "home. The Lambda first verifies that the requesting user holds either "
        "ADMIN or GUEST permissions on {home_id}, then writes a new row into "
        "automation_rules together with the associated rule_actions entries — "
        "each pinned to a specific (deviceID, property_name, value) triple. "
        "Sensor-triggered rules carry trigger_type: \"sensor\" with a numeric "
        "comparator (>, <, ==), while emotion-triggered rules carry "
        "trigger_type: \"emotion\" with a single mood label drawn from the "
        "same ten-class taxonomy used by the AI module. The Flutter "
        "AutomationCreateScreen serialises its three-step wizard state into "
        "the body shape shown below and submits it through "
        "ApiService.saveAutomation.",
    )
    insert_paragraph_at(doc, edge_anchor,
                        "Request body example (turn lights warm orange when emotion is angry):")
    for line in (
        '{',
        '  "rule_name": "Calm Down Lights",',
        '  "trigger_type": "emotion",',
        '  "trigger_emotion": "angry",',
        '  "actions": [',
        '    { "deviceID": "a1b2c3d4-...", "property_name": "color",      "value": "#FF6B35" },',
        '    { "deviceID": "a1b2c3d4-...", "property_name": "brightness", "value": "30" }',
        '  ],',
        '  "is_active": true',
        '}',
    ):
        insert_paragraph_at(doc, edge_anchor, line, role="code")

    insert_paragraph_at(doc, edge_anchor,
                        "Request body example (sensor trigger — turn off the AC when temperature drops below 20 °C):")
    for line in (
        '{',
        '  "rule_name": "Save Energy",',
        '  "trigger_type": "sensor",',
        '  "trigger_field": "temperature",',
        '  "comparator": "<",',
        '  "threshold": 20,',
        '  "actions": [',
        '    { "deviceID": "9f8e7d6c-...", "property_name": "power", "value": "off" }',
        '  ],',
        '  "is_active": true',
        '}',
    ):
        insert_paragraph_at(doc, edge_anchor, line, role="code")

    # New H4 + body for deleteAutomation
    insert_paragraph_at(
        doc, edge_anchor,
        "DELETE /{home_id}/automations?rule_id={rule_id} — deleteAutomation",
        style="Heading 4",
    )
    insert_paragraph_at(
        doc, edge_anchor,
        "Removes the automation rule identified by rule_id (a UUIDv4 query "
        "parameter) from automation_rules and cascades the deletion across "
        "rule_actions through the foreign-key relationship. Authorisation "
        "requires the same ADMIN/GUEST membership check as the other "
        "home-scoped endpoints. The mobile UI invokes this endpoint when the "
        "user confirms a swipe-to-delete gesture on a rule card; the Lambda "
        "responds with HTTP 200 on a successful deletion or HTTP 401 when the "
        "caller lacks permission for the home. No request body is required.",
    )


# ---------------------------------------------------------------------------
# New-chapter content stream
# ---------------------------------------------------------------------------

SECTIONS = []


def H2(t): SECTIONS.append(("Heading 2", t))
def H3(t): SECTIONS.append(("Heading 3", t))
def H4(t): SECTIONS.append(("Heading 4", t))
def P(t):  SECTIONS.append(("Normal", t))
def CAPTION(t): SECTIONS.append(("Caption", t))


def TABLE(headers, rows):
    SECTIONS.append(("__TABLE__", (headers, rows)))


def CODE(text):
    for line in text.splitlines():
        SECTIONS.append(("__CODE__", line))


def CODE_LINE(line):
    SECTIONS.append(("__CODE__", line))


def FIG(number, description, caption, kind="screenshot"):
    SECTIONS.append(("__FIG__", (number, description, caption, kind)))


def IMG(number, image_path, caption):
    SECTIONS.append(("__IMG__", (number, image_path, caption)))


# ============================================================================
# 1. Mobile Application Implementation (Flutter)
# ============================================================================
H2("Mobile Application Implementation (Flutter)")
P(
    "The client-side application is implemented as a single Flutter codebase "
    "(Dart 3.10, Flutter 3.38) targeting both Android and iOS from the same "
    "source tree. The code is organised under lib/ into five strictly separated "
    "layers — services/ (network and SDK adapters), providers/ (Riverpod state "
    "containers), screens/ (user-facing pages), widgets/ (reusable composable "
    "UI elements), and theme/ together with constants/ (visual tokens). This "
    "separation enforces a unidirectional data flow in which screens never "
    "invoke network APIs directly; they read from providers and dispatch user "
    "intents back through them."
)

H3("Application Bootstrap and Lifecycle")
P(
    "Application start-up is centralised in main.dart. Before runApp is invoked, "
    "three asynchronous initialisations are performed sequentially: dotenv.load() "
    "reads the bundled .env file containing the Claude, Spotify, and Pi endpoint "
    "secrets; Firebase.initializeApp() establishes the Firebase Cloud Messaging "
    "context required for push notifications; and Amplify.configure() wires the "
    "Cognito Auth plugin using the configuration generated by AWS Amplify. Once "
    "these complete, the app is wrapped in a ProviderScope widget so that the "
    "entire widget tree shares a single Riverpod container, and a "
    "GlobalKey<NavigatorState> is registered to allow background isolates (such "
    "as the FCM handler) to push dialogs without holding a BuildContext."
)

H3("State Management with Riverpod")
P(
    "State is managed through flutter_riverpod 3.x, in which seven NotifierProvider "
    "instances are declared. Table 5.5 summarises their responsibilities."
)
CAPTION("Table 5.5: Riverpod providers and their responsibilities")
TABLE(
    ["Provider", "Type", "Responsibility"],
    [
        ["authProvider", "AuthState enum", "Mirrors the Cognito session lifecycle and drives the root navigator (initial / loading / authenticated / unauthenticated)."],
        ["homeProvider", "Home (UUID + role)", "Holds the currently selected home; rebuilt every time the user switches."],
        ["moodProvider", "MoodState (mood, confidence, updatedAt, source)", "Single source of truth shared by Emotion Hub, AI chat, and Spotify."],
        ["alertProvider", "List<AlertItem>", "Persistent in-memory feed of FCM alerts."],
        ["themeProvider", "ThemeMode", "Light/dark toggle, persisted via shared_preferences."],
        ["spotifyProvider", "SpotifyAuthState", "Exposes isAuthenticated and login progress to UI."],
        ["navigationProvider", "int", "Selected tab index for the bottom navigation shell."],
    ],
)
P(
    "The moodProvider plays a central role in the architecture. When the Emotion "
    "Hub completes a face scan, when the user selects a manual mood pill, or when "
    "the AI chat agent invokes its set_mood tool, all three writers funnel "
    "through the same MoodNotifier.set(mood, confidence, source: ...) mutation. "
    "Listeners in the Spotify screen react via ref.listen(moodProvider, ...), "
    "automatically refetching personalised tracks whenever the underlying mood "
    "changes — without any imperative wiring between the three writers."
)

H3("Screen Hierarchy and Navigation")
P(
    "The top-level navigation is determined by authState: while the value is "
    "initial or loading, a centred CircularProgressIndicator is shown; on "
    "unauthenticated, the user lands on OnboardingScreen; on authenticated, the "
    "user is routed to HomeSelectionScreen. After a home is selected, the user "
    "enters MainShell, a persistent bottom-navigation scaffold that hosts the "
    "Dashboard, Devices, Automations, AI Hub, and Profile tabs. Modal flows such "
    "as the AI chat (AiChatScreen), Spotify (SpotifyTestScreen), notification "
    "list (NotificationScreen), security monitoring (MonitoringScreen), and the "
    "QR-based home invitation are pushed on top of this shell as full-screen "
    "routes."
)
IMG(
    "5.10",
    "images/fig_5_10_app_screen_flow.png",
    "Application screen flow and navigation graph.",
)

H3("Authentication Flow (Amplify + Cognito)")
P(
    "User authentication is fully delegated to AWS Amplify's Cognito plugin. "
    "Five screens implement the complete identity lifecycle: LoginScreen, "
    "RegisterScreen, OtpVerificationScreen, ForgotPasswordScreen, and "
    "ResetConfirmationScreen. Sign-in invokes Amplify.Auth.signIn(username, "
    "password); sign-up invokes Amplify.Auth.signUp(...) followed by "
    "confirmSignUp(username, code) once the user submits the six-digit OTP "
    "delivered to their inbox. The password reset flow follows the same two-step "
    "pattern using resetPassword and confirmResetPassword. After every "
    "successful authentication event, AuthProvider calls "
    "Amplify.Auth.fetchAuthSession() to obtain a fresh ID Token, the raw form of "
    "which is then attached as Authorization: Bearer <token> to every subsequent "
    "REST call by ApiService._getHeaders(). No password or token is persisted "
    "manually by the application — Amplify handles secure on-device storage "
    "transparently."
)
FIG(
    "5.11",
    "Side-by-side phone screenshots of LoginScreen (left), RegisterScreen "
    "(middle), and OtpVerificationScreen (right) showing the full sign-up "
    "path with the six-digit Cognito OTP entry.",
    "Authentication screens: login, register, and OTP verification.",
)

H3("Network Layer (ApiService) — Request and Response Payload Contracts")
P(
    "All AWS REST traffic is funnelled through a single ApiService class that "
    "points to the production base URL "
    "(https://zz3kr12z0f.execute-api.us-east-1.amazonaws.com/prod). Each method "
    "obtains the Cognito ID Token via Amplify.Auth.fetchAuthSession(), attaches "
    "it as Authorization: Bearer <token>, performs the HTTP call with a 15-second "
    "timeout, and parses the response. A network-suppression mechanism is "
    "implemented to prevent log flooding during intermittent connectivity "
    "loss: a static _networkDown flag is raised on the first SocketException "
    "and only reset once a subsequent request succeeds, after which a single "
    "Network recovered line is logged."
)
P(
    "The exact request and response payload contracts of every ApiService method "
    "are documented below. These bodies are the canonical wire format and are "
    "shared verbatim between the dashboard, device control, automations, and "
    "AI chat agent paths."
)

H4("sendCommand — POST /prod/{home_id}/command")
P(
    "Sends a single property change to one actuator. The mobile app and the AI "
    "chat agent both invoke this method; the agent calls it once per attribute "
    "when batching multi-property updates (e.g. a power toggle followed by a "
    "brightness change is dispatched as two separate calls). The body always "
    "wraps a single-element commands array — the schema permits multiple "
    "entries but the current implementation issues one property change per "
    "request to keep failure isolation simple."
)
P("Request body example (set LED brightness to 75):")
for line in (
    '{',
    '  "deviceID": "a1b2c3d4-e5f6-7890-abcd-ef0123456789",',
    '  "commands": [',
    '    { "property_name": "brightness", "value": "75" }',
    '  ]',
    '}',
):
    CODE_LINE(line)
P("Response: HTTP 200 on successful publish to IoT Core; HTTP 4xx if the deviceID does not belong to the home or the property is unknown for the device type.")

H4("fetchSensors — GET /prod/{home_id}/sensor")
P(
    "Pulled by the dashboard every five seconds. The Lambda reads the latest "
    "row per (deviceID, sensor_type) pair from sensor_data and returns a flat "
    "snapshot keyed by deviceID. Stale-data dimming on the dashboard is driven "
    "by the last_updated timestamp returned for each sensor."
)
P("Response body example:")
for line in (
    '{',
    '  "homeID": "e4b3c7d2-...",',
    '  "sensors": {',
    '    "9f8e7d6c-...": {',
    '      "temperature":  "24.5",',
    '      "humidity":     "45",',
    '      "last_updated": "2026-04-28T17:50:00Z"',
    '    },',
    '    "1a2b3c4d-...": {',
    '      "gas_level":    "300",',
    '      "status":       "safe",',
    '      "last_updated": "2026-04-28T17:50:05Z"',
    '    }',
    '  }',
    '}',
):
    CODE_LINE(line)

H4("fetchDevices — GET /prod/{home_id}/devices")
P(
    "Returns every actuator registered for the home, including its current "
    "operational state pulled from the actuator_current_states table. Each "
    "device entry carries a polymorphic properties array whose shape depends "
    "on the device_type — speakers add volume / playback, LEDs add color / "
    "brightness, blinds add position, and so on. The AI chat agent always "
    "calls this endpoint before invoking control_device so that it can "
    "reference the canonical deviceid string rather than a human-readable name."
)
P("Response body example:")
for line in (
    '{',
    '  "devices": [',
    '    {',
    '      "deviceid":    "a1b2c3d4-...",',
    '      "device_name": "Living Room LED",',
    '      "device_type": "lighting",',
    '      "properties": [',
    '        { "property_name": "power",      "value": "on"      },',
    '        { "property_name": "brightness", "value": "75"      },',
    '        { "property_name": "color",      "value": "#FFAA00" }',
    '      ]',
    '    },',
    '    {',
    '      "deviceid":    "b2c3d4e5-...",',
    '      "device_name": "Bedroom Blinds",',
    '      "device_type": "curtain",',
    '      "properties": [',
    '        { "property_name": "position", "value": "100" }',
    '      ]',
    '    }',
    '  ]',
    '}',
):
    CODE_LINE(line)

H4("fetchAutomations — GET /prod/{home_id}/automations")
P(
    "Lists every automation rule configured for the home. The Lambda joins "
    "automation_rules with rule_actions and devices so that each returned rule "
    "carries enough context to be rendered as a single card in the "
    "AutomationsListScreen without follow-up calls."
)
P("Response body example:")
for line in (
    '{',
    '  "automations": [',
    '    {',
    '      "rule_id":         "r-7890abcd-...",',
    '      "rule_name":       "Calm Down Lights",',
    '      "trigger_type":    "emotion",',
    '      "trigger_emotion": "angry",',
    '      "is_active":       true,',
    '      "actions": [',
    '        { "device_name": "Living Room LED", "property_name": "color",      "value": "#FF6B35" },',
    '        { "device_name": "Living Room LED", "property_name": "brightness", "value": "30"      }',
    '      ]',
    '    }',
    '  ]',
    '}',
):
    CODE_LINE(line)

H4("saveAutomation — POST /prod/{home_id}/automations")
P(
    "Documented in the existing chapter alongside the other home-scoped "
    "endpoints. The Flutter AutomationCreateScreen serialises its three-step "
    "wizard state into either the sensor-trigger or the emotion-trigger body "
    "shape and submits it directly through ApiService.saveAutomation."
)

H4("deleteAutomation — DELETE /prod/{home_id}/automations?rule_id={rule_id}")
P(
    "Invoked when the user confirms a swipe-to-delete gesture on a rule card. "
    "No request body is required; the rule UUID travels as a query string "
    "parameter so the deletion is idempotent and cache-friendly."
)

H4("saveFcmToken — PUT /prod/fcm-token")
P(
    "Called once on every app launch and again whenever FlutterFire's "
    "onTokenRefresh fires. The body carries only the FCM registration token; "
    "the Lambda derives the user UUID from the Cognito claims so the request "
    "remains tamper-resistant."
)
P("Request body example:")
for line in (
    '{',
    '  "fcm_token": "fXyZ3pE2RQabc...:APA91bH...long-fcm-token..."',
    '}',
):
    CODE_LINE(line)

H3("Dashboard and Sensor Polling")
P(
    "The dashboard (DashboardScreen) displays live environmental telemetry. "
    "Rather than maintaining a persistent WebSocket connection, the screen "
    "executes a lightweight five-second polling loop against fetchSensors. "
    "The returned snapshot is rendered through the SensorTile widget, which "
    "colour-codes thresholds (gas above warning level turns the tile orange, "
    "above danger turns it red). Stale-data detection is based on the "
    "last_updated timestamp: tiles older than thirty seconds are visually "
    "dimmed to alert the user that the corresponding sensor may be offline. "
    "The dashboard widgets were extracted into lib/screens/dashboard/widgets/ "
    "during the refactoring sprint to enable independent testing and reuse."
)
FIG(
    "5.12",
    "Phone screenshot of DashboardScreen showing live sensor tiles for "
    "temperature, humidity, gas concentration, and vibration intensity, with "
    "threshold-driven colour coding (green / orange / red).",
    "Dashboard screen with live sensor tiles.",
)

H3("Device Control Screen")
P(
    "DeviceControlScreen renders each actuator returned by fetchDevices as an "
    "interactive card. Device-type icons are inferred polymorphically: the "
    "system first inspects the device_type column from the database, and falls "
    "back to keyword matching against the device_name field (covering Turkish "
    "and English variants such as door / kapı, speaker / hoparlör, stove / "
    "fırın, window / pencere, outlet / priz). Property writes are dispatched "
    "through ApiService.sendCommand, whose canonical body shape is documented "
    "in the previous sub-section. Optimistic UI updates are applied "
    "immediately; the underlying state is reconciled when updateActuatorState "
    "writes the actual physical state back through the digital-twin path."
)
FIG(
    "5.13",
    "Three control cards from DeviceControlScreen: an RGB LED strip with "
    "colour picker and brightness slider, a speaker with volume and playback "
    "controls, and a curtain motor with position slider.",
    "Device control screen examples (lighting, speaker, curtain).",
)

H3("Automations List and Creation")
P(
    "The automations module consists of two screens: AutomationsListScreen and "
    "AutomationCreateScreen. The list screen displays a custom header with live "
    "counters (Total / Active / AI), filter chips (All / Active / AI / Sensor), "
    "and colour-accented cards (orange for emotion-driven rules, blue for "
    "sensor-driven rules). Swipe-to-delete with a confirmation dialog and "
    "pull-to-refresh are implemented through Dismissible and RefreshIndicator "
    "respectively. The creation screen guides the user through a three-step "
    "wizard: trigger type selection (sensor or mood), trigger parameter "
    "configuration (e.g., temperature > 30°C or emotion = happy), and "
    "multi-device action selection (each action specifies a target device, "
    "property, and value). The completed rule is serialised to the JSON body "
    "shape documented under saveAutomation in the existing chapter and "
    "submitted via POST /prod/{home_id}/automations."
)
FIG(
    "5.14",
    "AutomationsListScreen (left) showing filter chips, live counters and "
    "colour-coded cards for AI- and sensor-driven rules; "
    "AutomationCreateScreen (right) showing the three-step wizard.",
    "Automations list and creation wizard.",
)


# ============================================================================
# 2. Edge Layer Implementation (Raspberry Pi)
# ============================================================================
H2("Edge Layer Implementation (Raspberry Pi)")
P(
    "The Raspberry Pi 5 acts as the on-premise intelligence node and hardware "
    "bridge. Three concurrent processes run at all times: an MQTT client process "
    "that maintains the bidirectional channel with AWS IoT Core, a "
    "sensor-polling daemon that publishes telemetry to the appropriate topics, "
    "and a FastAPI service hosting the /predict endpoint for facial emotion "
    "analysis. All three are supervised by systemd so that a crash in any single "
    "process does not require manual intervention."
)
FIG(
    "5.15",
    "Photograph of the assembled Raspberry Pi 5 prototype: the Pi board "
    "mounted on a breadboard, with DHT11, MQ-3 / MQ-4, MPU6050, LDR and RFID "
    "sensors wired in, alongside the LED strip, stepper-driven curtain motor, "
    "and safety siren.",
    "Physical Raspberry Pi prototype with sensors and actuators.",
    kind="photo",
)

H3("MQTT Client and Topic Subscription")
P(
    "The MQTT layer is built on paho-mqtt. On startup, the client establishes a "
    "TLS 1.2 connection to AWS IoT Core using the X.509 device certificate "
    "provisioned during the IoT Thing registration shown earlier in Figure 5.6. "
    "Two subscriptions are registered: home/{home_id}/commands for incoming "
    "actuator instructions and a system-level wildcard for diagnostic topics. "
    "The on-message callback dispatches each payload through a lightweight "
    "router that inspects the deviceID field and forwards the command list to "
    "the corresponding hardware driver. After physical execution, the same "
    "client publishes a state acknowledgement to home/{home_id}/state with "
    "QoS 1, ensuring at-least-once delivery so that the cloud digital twin is "
    "never permanently out of sync with reality."
)

H3("Hardware Drivers (GPIO, PWM, I²C)")
P(
    "Each actuator type is encapsulated in a dedicated Python module that "
    "abstracts the underlying GPIO interface from the MQTT router. Lighting "
    "(LED Strip) is driven via PWM on three colour channels (red, green, blue) "
    "for full RGB control: the brightness property scales the duty cycle "
    "uniformly across all channels and the color property parses #RRGGBB hex "
    "strings into per-channel duty cycles. The curtain / blinds motor is a "
    "stepper motor controlled through GPIO sequencing — the position property "
    "(0–100) is mapped to a target step count, and the motor is advanced or "
    "retracted from the last known position with software limit enforcement. "
    "The climate module (AC / Stove) uses boolean relay channels controlled "
    "through digital-write on dedicated GPIO pins. The speaker is driven through "
    "PWM amplifier control for volume, with playback commands (play/pause/stop) "
    "routed through pygame.mixer for local audio playback. Finally, the safety "
    "siren is a high-priority GPIO output activated locally by the "
    "hazard-detection daemon before any network call, ensuring the alarm sounds "
    "even during cloud or internet outages."
)
IMG(
    "5.16",
    "images/fig_5_16_pi_wiring.png",
    "Sensor and actuator wiring overview on the Raspberry Pi 5.",
)

H3("Sensor Polling and Anomaly Detection")
P(
    "The sensor daemon runs an event loop with three timing classes. Periodic "
    "sensors (DHT11, LDR) are polled every 10 seconds, with the temperature, "
    "humidity, and light-intensity readings published to home/{home_id}/sensor "
    "as a single bundled JSON payload. Interrupt-driven sensors (MPU6050 "
    "vibration, MQ-3 / MQ-4 gas) are configured as edge-triggered GPIO "
    "interrupts: when a threshold is crossed, an immediate publish to "
    "home/{home_id}/alert is performed before any further processing — "
    "minimising the time between physical event and user notification. "
    "On-demand peripherals (RFID reader, micro camera) are idle by default and "
    "activated only when the cloud explicitly requests a read or when the "
    "FastAPI emotion service is invoked."
)
P(
    "When the gas concentration crosses the danger threshold or the MPU6050 "
    "reports vibration intensity above the seismic-event threshold, two parallel "
    "actions are taken: the local siren GPIO is asserted immediately for "
    "autonomous safety response, and the alert payload defined earlier in this "
    "chapter is published to AWS IoT Core, which in turn invokes the "
    "pushAlertNotificationsToUsers Lambda. The combination of edge-side "
    "actuation and cloud-side notification guarantees that residents are warned "
    "both audibly inside the home and remotely on their mobile devices."
)

H3("Facial Emotion Recognition Service — POST /predict")
P(
    "The Raspberry Pi hosts a FastAPI application listening on "
    "https://<pi_host>:8000/predict (defaulting to https://ramazan.local:8000/"
    "predict via mDNS). The endpoint accepts a multipart-encoded JPEG image "
    "and returns a JSON response containing the dominant emotion label, its "
    "confidence score, and the full per-class softmax distribution."
)
P("Request format (multipart/form-data; the only required part is file):")
for line in (
    'POST /predict HTTP/1.1',
    'Host: ramazan.local:8000',
    'Content-Type: multipart/form-data; boundary=----flutter-boundary',
    '',
    '------flutter-boundary',
    'Content-Disposition: form-data; name="file"; filename="frame.jpg"',
    'Content-Type: image/jpeg',
    '',
    '<binary JPEG bytes>',
    '------flutter-boundary--',
):
    CODE_LINE(line)
P("Response body example (success path):")
for line in (
    '{',
    '  "status": "success",',
    '  "emotion": "happy",',
    '  "confidence": 0.8722,',
    '  "all_scores": {',
    '    "angry": 0.01, "disgusted": 0.00, "fearful": 0.02,',
    '    "happy": 0.87, "neutral": 0.05, "sad": 0.03, "surprised": 0.02',
    '  }',
    '}',
):
    CODE_LINE(line)
P("Response body example (no face detected):")
for line in (
    '{',
    '  "status": "no_face_detected",',
    '  "emotion": null,',
    '  "confidence": 0.0',
    '}',
):
    CODE_LINE(line)
P(
    "The processing pipeline is performed entirely on-device. The incoming JPEG "
    "is loaded via OpenCV (cv2.imdecode) into a BGR numpy array. A Haar cascade "
    "classifier then detects the dominant face region and crops it; if no face "
    "is found, the no_face_detected status is returned and the mobile app "
    "surfaces a user-friendly toast. The cropped region is converted to "
    "greyscale, resized to 48×48 pixels, and normalised to the [0, 1] range — "
    "matching the input shape of the trained Keras model. The pre-processed "
    "tensor is fed through a TensorFlow Lite model derived from a CNN trained "
    "on the FER-2013 dataset, which outputs a softmax distribution over seven "
    "emotion classes. The argmax class is reported as emotion, its softmax "
    "probability as confidence, and the full distribution as all_scores so that "
    "the mobile UI can render confidence bars per class."
)
P(
    "Because the model runs locally and the JPEG is never persisted to disk, no "
    "facial imagery ever leaves the home network. This satisfies the KVKK and "
    "GDPR commitments stated in Chapter 1 and the edge-processing constraints "
    "defined in Chapter 3. The TLS layer uses a self-signed certificate "
    "provisioned at the time of Pi setup; in production deployments this is "
    "replaced by a Let's Encrypt certificate routed via mDNS (ramazan.local). "
    "The standalone helper script camera.py at the repository root exercises the "
    "same endpoint from a desktop webcam using OpenCV and requests, providing a "
    "convenient way to validate the model end-to-end without launching the "
    "mobile application."
)


# ============================================================================
# 3. AI-Driven Emotion Recognition Module (Mobile Side)
# ============================================================================
H2("AI-Driven Emotion Recognition Module (Mobile Side)")
P(
    "The mobile counterpart of the emotion recognition pipeline is implemented "
    "in EmotionApiService and consumed by the Emotion Hub screen. Capture, "
    "transport, and rendering of the result are all performed in under two "
    "seconds on a typical Wi-Fi network."
)

H3("Image Capture and Transmission")
P(
    "The Emotion Hub screen exposes a single circular Scan button. When tapped, "
    "image_picker.pickImage(source: ImageSource.camera, preferredCameraDevice: "
    "CameraDevice.front) opens the device front camera. The user takes a photo, "
    "which is returned as a JPEG file. The bytes are read via "
    "File.readAsBytes() and passed to "
    "EmotionApiService.predictFromJpeg(jpegBytes). The service constructs an "
    "http.MultipartRequest with a single file part named frame.jpg and POSTs it "
    "to the configured EMOTION_API_URL — the wire format is documented in "
    "the previous Pi sub-section. A 15-second timeout is enforced to avoid "
    "indefinite UI blocking on Pi unavailability."
)

H3("Self-Signed Certificate Handling")
P(
    "Because the Pi serves the endpoint over HTTPS using a self-signed "
    "certificate, the standard Dart HTTP client would reject the connection. To "
    "avoid this during development, _buildClient() returns an IOClient wrapping "
    "an HttpClient whose badCertificateCallback is set to always accept — but "
    "only when kDebugMode is true. Release builds fall back to a plain "
    "http.Client() that enforces standard certificate validation. The expected "
    "production deployment installs a Let's Encrypt or private-CA certificate on "
    "the Pi, after which the bypass becomes a no-op."
)

H3("Result Parsing and UI Rendering")
P(
    "The service tolerates two possible response shapes (emotion or mood as the "
    "label key, confidence or score as the score key), wraps the result in an "
    "EmotionApiResult value object, and returns it to the calling screen. The "
    "Emotion Hub then writes the result into moodProvider with source: 'scan'. "
    "A colour and emoji are looked up from internal maps that cover all ten "
    "supported mood classes (happy, sad, melancholy, angry, calm, excited, "
    "neutral, fearful, surprised, disgusted) — the scan ring animates from "
    "neutral grey to the mood-specific colour, the confidence percentage is "
    "displayed numerically, and an Ambient Suggestion tile recommends a lighting "
    "tone and curtain position appropriate to the detected emotion."
)
FIG(
    "5.17",
    "Emotion Hub before scan (left, neutral grey scan ring) and after a "
    "successful scan (right, mood-coloured ring with emoji, confidence "
    "percentage, and an Ambient Suggestion tile recommending a lighting "
    "tone and curtain position).",
    "Emotion Hub screen — scan ring states and ambient suggestion tile.",
)

H3("Manual Mood Picker (Fallback)")
P(
    "When the camera is unavailable (e.g., on the Android emulator, where the "
    "simulated frame contains no real face) or the user disagrees with the "
    "model's interpretation, a manual fallback is provided. Tapping a secondary "
    "Manual button opens a bottom sheet containing ten tappable mood pills. "
    "Selecting a pill writes the chosen mood into moodProvider with a confidence "
    "of 1.0 and source: 'manual'. Downstream consumers (Spotify, automations, "
    "AI chat) treat manual entries identically to scan results."
)
FIG(
    "5.18",
    "Manual mood picker bottom sheet showing the ten tappable mood pills "
    "(happy, sad, melancholy, angry, calm, excited, neutral, fearful, "
    "surprised, disgusted) with mood-specific colour fills.",
    "Manual mood picker bottom sheet.",
)


# ============================================================================
# 4. AI Chat Agent (Claude with Tool Use)
# ============================================================================
H2("AI Chat Agent (Claude with Tool Use)")
P(
    "The conversational assistant is implemented in AiAgentService and consumed "
    "by AiChatScreen. It is the most sophisticated AI component of the system, "
    "since it can reason over the user's current home state and execute control "
    "actions on their behalf via the Anthropic Tool Use API."
)

H3("Model Selection and System Prompt")
P(
    "The agent uses claude-haiku-4-5-20251001, selected for its low latency "
    "(sub-second time-to-first-token on most prompts) and strong tool-use "
    "reliability. The system prompt establishes nine non-negotiable rules: every "
    "device control must be preceded by a get_devices call to obtain the "
    "canonical deviceid; one property change is sent per control_device "
    "invocation; sensor readings must be reported with units; replies must match "
    "the user's language (Turkish or English); and any user disclosure of "
    "feeling — including corrections of a previous scan — must trigger set_mood."
)

H3("Tool Definitions")
P(
    "Five tools are registered with the Messages API. Their schemas are "
    "summarised in Table 5.6."
)
CAPTION("Table 5.6: Tools exposed to the Claude Haiku agent")
TABLE(
    ["Tool", "Inputs", "Effect"],
    [
        ["get_devices", "none", "Calls ApiService.fetchDevices(homeId) and returns the device list."],
        ["get_sensor_data", "none", "Calls ApiService.fetchSensors(homeId) and returns the latest snapshot."],
        ["get_automations", "none", "Calls ApiService.fetchAutomations(homeId) and returns active rules."],
        ["control_device", "device_id, action, value", "Invokes ApiService.sendCommand(...), which publishes through Lambda → IoT Core → Pi."],
        ["set_mood", "mood (enum), confidence (0.0–1.0)", "Updates moodProvider directly via the onSetMood callback."],
    ],
)
P(
    "The set_mood tool is unique in that it does not perform any network call: "
    "its handler is a closure passed by AiChatScreen that mutates the shared "
    "Riverpod state, and the resulting state change is immediately observable "
    "by the Emotion Hub and Spotify screens through ref.listen subscriptions."
)

H3("Conversation Loop and Anthropic API Request Format")
P(
    "AiAgentService.chat() implements an iteration-bounded loop (maximum five "
    "iterations, to prevent runaway tool chains). On each iteration the full "
    "message history is sent to https://api.anthropic.com/v1/messages along with "
    "the tool list. If the model returns stop_reason == \"tool_use\", the "
    "response is parsed into one or more tool_use blocks; each block is "
    "dispatched to the corresponding handler; the results are appended back to "
    "the conversation as a tool_result user message; and the loop continues. If "
    "the model returns stop_reason == \"end_turn\", the textual content is "
    "returned to the UI as the assistant reply."
)
P("Request body submitted to POST https://api.anthropic.com/v1/messages:")
for line in (
    'POST https://api.anthropic.com/v1/messages',
    'Headers:',
    '  Content-Type:       application/json',
    '  x-api-key:          <CLAUDE_API_KEY from .env>',
    '  anthropic-version:  2023-06-01',
    'Body:',
    '{',
    '  "model": "claude-haiku-4-5-20251001",',
    '  "max_tokens": 1024,',
    '  "system": "You are a helpful smart home AI assistant... (full system prompt)",',
    '  "messages": [',
    '    { "role": "user", "content": "lambaları kapat ve sakin müzik çal" }',
    '  ],',
    '  "tools": [',
    '    { "name": "get_devices",     "description": "...", "input_schema": { ... } },',
    '    { "name": "get_sensor_data", "description": "...", "input_schema": { ... } },',
    '    { "name": "get_automations", "description": "...", "input_schema": { ... } },',
    '    { "name": "control_device",  "description": "...", "input_schema": { ... } },',
    '    { "name": "set_mood",        "description": "...", "input_schema": { ... } }',
    '  ]',
    '}',
):
    CODE_LINE(line)
P("Tool-use response shape (when the model decides to call a tool):")
for line in (
    '{',
    '  "stop_reason": "tool_use",',
    '  "content": [',
    '    { "type": "text", "text": "Tamam, önce cihazları listeleyeyim." },',
    '    {',
    '      "type": "tool_use",',
    '      "id": "toolu_01AaBbCc...",',
    '      "name": "get_devices",',
    '      "input": {}',
    '    }',
    '  ]',
    '}',
):
    CODE_LINE(line)
P(
    "The service then issues fetchDevices, packages the JSON response as a "
    "tool_result user message keyed by the same id, and re-invokes the API. "
    "This loop allows multi-step plans such as \"turn off the lights, set the "
    "AC to off, and play calm music\" to be executed within a single user turn "
    "through three independent control_device invocations."
)
FIG(
    "5.19",
    "AiChatScreen showing a multi-step user request (\"lambaları kapat ve "
    "sakin müzik çal\") and the agent's response, with inline indicators "
    "marking the get_devices and control_device tool invocations performed "
    "during the turn.",
    "AI chat agent conversation with tool-use indicators.",
)

H3("Mood-Bound Conversational Flow")
P(
    "A representative end-to-end trace illustrates how the modules cooperate. "
    "The user sends the message \"Ben aslında üzgünüm, az önceki tarama yanlıştı\" "
    "in the chat. Claude parses the disclosure, calls "
    "set_mood(mood=\"sad\", confidence=1.0), which mutates moodProvider. The "
    "Emotion Hub, listening on moodProvider, redraws its scan ring in the "
    "sad-mood colour with a (manual) indicator. Concurrently, the Spotify "
    "screen — also subscribed to the same provider — re-runs its mood pipeline "
    "and replaces the displayed track list with sad-mood-matched personal "
    "favourites. The user has issued a single natural-language sentence, and "
    "three independent UI surfaces have updated coherently without any direct "
    "cross-screen coupling."
)
FIG(
    "5.20",
    "Three-panel composition: AiChatScreen with the user's mood-correction "
    "message (left), Emotion Hub redrawn in the sad-mood colour with a "
    "(manual) indicator (centre), and Spotify screen showing newly fetched "
    "sad-mood-matched personal tracks (right).",
    "Mood-bound conversational flow across chat, emotion hub, and Spotify.",
)


# ============================================================================
# 5. Spotify Integration
# ============================================================================
H2("Spotify Integration (OAuth, Mood-Driven Recommendations, Claude as Music Curator)")
P(
    "Spotify integration is implemented under lib/services/spotify/ as four "
    "cooperating modules: SpotifyAuth (OAuth lifecycle), SpotifyApiClient (raw "
    "REST wrappers), SpotifyMoodCatalog (keyword-based mood matching), and "
    "SpotifyClaudeAgent (LLM-driven candidate generation). The public "
    "SpotifyService class is a thin façade that screens and providers consume "
    "without ever touching the internal modules directly."
)

H3("OAuth 2.0 Authorisation Code Flow with PKCE")
P(
    "User authorisation is performed through a fully-fledged OAuth 2.0 "
    "Authorization Code flow. SpotifyAuth.login() constructs an authorisation "
    "URL targeting https://accounts.spotify.com/authorize with the registered "
    "client_id, the redirect URI akilliev://callback, and the six scopes the "
    "app requires (user-read-recently-played, user-top-read, "
    "user-read-playback-state, user-modify-playback-state, "
    "user-read-currently-playing, playlist-read-private). The URL is opened "
    "through flutter_web_auth_2, which presents an in-app browser session. Upon "
    "successful login the user is redirected back to the registered scheme and "
    "the captured code is exchanged for tokens via a POST to the Spotify "
    "Accounts API."
)
P("Token exchange request — POST https://accounts.spotify.com/api/token:")
for line in (
    'POST https://accounts.spotify.com/api/token',
    'Headers:',
    '  Content-Type:  application/x-www-form-urlencoded',
    '  Authorization: Basic <base64(client_id:client_secret)>',
    'Body (form-urlencoded):',
    '  grant_type=authorization_code',
    '  code=<auth_code_from_redirect>',
    '  redirect_uri=akilliev://callback',
):
    CODE_LINE(line)
P("Successful response body:")
for line in (
    '{',
    '  "access_token":  "BQDl0xY...",',
    '  "token_type":    "Bearer",',
    '  "expires_in":    3600,',
    '  "refresh_token": "AQDxyzABC...",',
    '  "scope":         "user-top-read user-read-recently-played ..."',
    '}',
):
    CODE_LINE(line)
P(
    "The access token, refresh token, and expiry timestamp are persisted in "
    "SharedPreferences so that subsequent app launches restore the session "
    "without prompting the user again. Tokens are refreshed automatically "
    "through a POST to the same URL with grant_type=refresh_token whenever the "
    "cached access token is within sixty seconds of expiry."
)
FIG(
    "5.21",
    "OAuth 2.0 PKCE flow for Spotify: in-app browser presenting the Spotify "
    "authorise prompt (left), permission scope confirmation (middle), and "
    "successful return to the application with the authenticated state "
    "indicator (right).",
    "Spotify OAuth 2.0 PKCE login flow.",
)

H3("Spotify Web API Endpoints Consumed by the App")
P(
    "Once authenticated, the app talks to the Spotify Web API through five "
    "GET endpoints. Every request carries Authorization: Bearer <access_token> "
    "in its headers; a 401 response triggers a transparent refresh-token "
    "exchange, after which the failed call is retried once. Table 5.7 "
    "summarises the call surface."
)
CAPTION("Table 5.7: Spotify Web API endpoints consumed by the application")
TABLE(
    ["Method", "URL", "Purpose"],
    [
        ["GET", "https://api.spotify.com/v1/me/top/tracks?limit=50&time_range=<short|medium|long>_term", "Pulls the user's top tracks for one of the three time horizons; called three times to build the personal pool."],
        ["GET", "https://api.spotify.com/v1/me/top/artists?limit=10&time_range=medium_term", "Pulls the user's top artists; feeds the Claude curator's profile object."],
        ["GET", "https://api.spotify.com/v1/me/player/recently-played?limit=20", "Pulls the most recently played tracks to enrich the listening profile."],
        ["GET", "https://api.spotify.com/v1/search?q=track:\"<title>\" artist:\"<name>\"&type=track&limit=1", "Resolves each Claude candidate to a real Spotify track."],
        ["GET", "https://api.spotify.com/v1/artists/{id}/top-tracks?market=TR", "Optional secondary path; deprecated for newly created apps and returns 403, retained for completeness."],
    ],
)
P(
    "The two `/me/top` endpoints return Spotify's standard paged track / artist "
    "objects, from which the app extracts only id, name, primary artist, album, "
    "popularity, uri, and (for artists) genres. The /search endpoint is used "
    "with quoted field filters (track: and artist:) so that title-only matches "
    "across different artists do not pollute the result."
)

H3("Working Around the November 2024 Spotify API Deprecation")
P(
    "In November 2024 Spotify removed /v1/recommendations, /v1/audio-features, "
    "/v1/audio-analysis, and /v1/artists/{id}/top-tracks from the public API "
    "for newly created applications. Since the original system design relied on "
    "the audio-features endpoint to infer mood compatibility, an alternative "
    "pipeline had to be engineered. The replacement is grounded entirely in "
    "endpoints that remain accessible (/v1/me/top/tracks and "
    "/v1/me/player/recently-played) and operates exclusively over the user's "
    "own listening history."
)

H3("Personalised Catalog Pipeline (SpotifyMoodCatalog)")
P(
    "The fallback recommendation pipeline pulls roughly 100–150 of the user's "
    "top tracks across three time ranges (short_term, medium_term, long_term) "
    "and merges them into a deduplicated personal pool. For the active mood, "
    "two parallel keyword lists — one English and one Turkish — are applied "
    "against each track's name, album title, and primary artist using "
    "word-boundary regular-expression matching (so that hüzün, kırık, yalnız, "
    "sad, broken, lonely and similar tokens flag relevant tracks). Each track "
    "receives a relevance score equal to the number of distinct keyword hits; "
    "tracks with score > 0 are tagged catalog_mood_matched and the remaining "
    "slots are filled from the unmatched personal pool with the tag "
    "catalog_fill. Because the candidate set is restricted to songs the user "
    "has actually listened to, a Turkish listener never receives random foreign "
    "pop-rock recommendations — every track is implicitly culturally and "
    "linguistically aligned."
)

H3("Claude Haiku as a Music Curator (SpotifyClaudeAgent)")
P(
    "When a Claude API key is configured, the agent-based path takes precedence "
    "over the keyword catalog. The agent operates in three stages. In the "
    "profile-construction stage, the user's top artists, top tracks "
    "(genre-deduplicated samples), and listening genres are aggregated into a "
    "compact profile object. In the candidate-generation stage, Claude Haiku is "
    "prompted with the profile and the active mood through the same Anthropic "
    "Messages endpoint described earlier, but with no tools — the model is "
    "asked to propose twenty real-world song candidates as {title, artist} "
    "pairs that fit both the user's taste and the requested emotional valence. "
    "The model is instructed not to repeat any title already in knownTitles to "
    "avoid recommending a song the user is currently listening to. In the "
    "catalog-resolution stage, each {title, artist} candidate is resolved to a "
    "concrete Spotify track through /v1/search, post-filtered for artist "
    "diversity (no more than two consecutive tracks by the same artist) and "
    "exclusion of already-known titles. Successfully resolved tracks are "
    "returned to the UI; unresolved candidates are silently dropped."
)
P(
    "A ten-minute in-memory TTL cache is keyed on the lowercased mood, "
    "eliminating redundant Claude calls during rapid mood changes within a "
    "single session. If the agent fails (missing API key, network error, empty "
    "candidate list, or zero successful resolutions), the system gracefully "
    "falls back to the catalog pipeline so that a song list is always returned "
    "to the UI."
)

H3("Mood-Reactive UI")
P(
    "The Spotify screen subscribes to moodProvider via ref.listen. Whenever the "
    "active mood changes — whether due to a face scan, a manual pick, or a "
    "set_mood invocation by the AI chat — the screen automatically re-runs the "
    "recommendation pipeline and replaces the displayed track list. Each track "
    "card shows the cover artwork, title, primary artist, and a small chip "
    "indicating its provenance (claude_agent, catalog_mood_matched, or "
    "catalog_fill), giving the user transparency into why each song was "
    "recommended."
)
FIG(
    "5.22",
    "Spotify mood recommendations screen with mood = happy: vertically "
    "scrolling list of personalised tracks, each card showing cover "
    "artwork, title, primary artist and a provenance chip "
    "(claude_agent / catalog_mood_matched / catalog_fill).",
    "Spotify mood-reactive recommendations screen.",
)


# ============================================================================
# 6. Push Notification System (FCM)
# ============================================================================
H2("Push Notification System (Firebase Cloud Messaging)")
P(
    "Push notifications are implemented end-to-end across three layers: token "
    "acquisition on the device, token persistence in the cloud, and notification "
    "delivery from the backend. The implementation is structured to deliver "
    "critical safety alerts (gas leaks, seismic events) within milliseconds of "
    "the edge detection, regardless of whether the app is in the foreground, "
    "background, or terminated."
)

H3("Token Lifecycle and Persistence Endpoint")
P(
    "On every app launch, the Firebase SDK either retrieves the existing FCM "
    "registration token or generates a new one if it is the first run on the "
    "device. The token is then forwarded to the cloud through the saveFcmToken "
    "endpoint documented earlier in the Network Layer section — the request "
    "carries only the FCM token in its body and the Lambda derives the user's "
    "UUID from the Cognito ID Token claims. Because the FCM token can be "
    "rotated by the operating system at any time (e.g., after an app data clear "
    "or a long inactivity period), the FlutterFire onTokenRefresh callback is "
    "wired to repeat the same persistence call automatically, ensuring the "
    "cloud always holds the device's currently valid routing token."
)

H3("Foreground Notification Handling")
P(
    "When a notification arrives while the app is open, "
    "FirebaseMessaging.onMessage.listen fires a callback in main.dart. The "
    "callback first extracts the title and body fields, gracefully tolerating "
    "both notification-shaped and data-shaped payloads (since the cloud sends "
    "data-only messages so that the OS does not auto-render a system tray "
    "banner that the in-app modal would then duplicate). The event type is "
    "mapped to an AlertType / AlertLevel pair: gas_leak and earthquake produce "
    "security / critical, all other events default to device / info. An "
    "AlertItem is appended to alertListProvider so it appears in the persistent "
    "in-app feed, and a non-dismissible red AlertDialog is presented through "
    "the global navigatorKey, ensuring the dialog is rendered even when no "
    "BuildContext is available to the caller."
)
P("Example FCM data payload received by the foreground handler:")
for line in (
    '{',
    '  "data": {',
    '    "title":     "EMERGENCY: Gas Leak Detected!",',
    '    "body":      "Dangerous gas levels detected in the kitchen.",',
    '    "event":     "gas_leak",',
    '    "severity":  "critical",',
    '    "deviceID":  "1a2b3c4d-...",',
    '    "homeID":    "e4b3c7d2-...",',
    '    "timestamp": "2026-04-28T19:40:00Z"',
    '  }',
    '}',
):
    CODE_LINE(line)
FIG(
    "5.23",
    "Foreground emergency alert: red non-dismissible AlertDialog raised on "
    "a gas-leak FCM message (left); NotificationScreen showing the "
    "chronological alert feed with colour-coded badges for critical, "
    "warning, and info levels (right).",
    "Foreground FCM alert dialog and persistent notification feed.",
)

H3("Background Notification Handling")
P(
    "When the app is in the background or terminated, a top-level function "
    "annotated with @pragma('vm:entry-point') is invoked by FlutterFire in a "
    "separate isolate. Because each Dart isolate has an independent memory "
    "space, Firebase must be re-initialised inside the handler before any "
    "plugin call. The handler currently logs the incoming notification for "
    "diagnostics; the visible notification banner itself is rendered "
    "automatically by the operating system from the data payload's "
    "notification fields, so no additional UI work is required from Dart in "
    "this path."
)

H3("Alert Feed and Persistence")
P(
    "The alertListProvider keeps a chronologically ordered list of every "
    "AlertItem received during the session. The NotificationScreen renders this "
    "list with colour-coded badges (red for critical, orange for warning, blue "
    "for info) and a relative timestamp. The list is intentionally "
    "session-scoped: critical alerts are also written to the cloud database by "
    "the pushAlertNotificationsToUsers Lambda, so the authoritative history is "
    "available server-side and can be retrieved via a future endpoint if "
    "required."
)


# ============================================================================
# 7. Theme System and Reusable UI Components
# ============================================================================
H2("Theme System and Reusable UI Components")
P(
    "A consistent visual identity is enforced through three coupled artefacts: "
    "app_theme.dart, app_colors.dart, and a small library of reusable widgets "
    "in lib/widgets/."
)

H3("Light and Dark Material 3 Themes")
P(
    "Two ThemeData instances are declared in AppTheme. Both seed from the same "
    "primaryBlue colour token but differ in surface, on-surface, and elevation "
    "values to produce coherent light and dark variants of every screen. The "
    "active theme is selected by themeProvider, the value of which is persisted "
    "in SharedPreferences so that the user's preference survives across "
    "launches. All colours referenced by widgets are routed through "
    "AppColors.bg(context), AppColors.fg(context), and similar helpers that "
    "resolve the correct shade based on the current Theme.of(context).brightness, "
    "so no widget hard-codes a literal colour."
)
FIG(
    "5.24",
    "Side-by-side comparison of the same DashboardScreen rendered under "
    "the light theme (left) and the dark theme (right), showing identical "
    "layout with theme-driven background, foreground, and accent colours.",
    "Light and dark theme comparison on the dashboard screen.",
)

H3("Mood-Aware Colour Palette")
P(
    "The Emotion Hub, Spotify card backgrounds, and the AI chat header all "
    "respect the same mood-to-colour mapping. Happy maps to a warm yellow, sad "
    "to a desaturated blue, calm to a pale green, angry to a saturated red, and "
    "so on for the full ten-class set. Because every consumer reads from this "
    "single map, an update to the palette propagates everywhere without any "
    "per-screen edit."
)
FIG(
    "5.25",
    "Swatch grid showing the ten mood-aware colour tokens (happy, sad, "
    "melancholy, angry, calm, excited, neutral, fearful, surprised, "
    "disgusted) together with the matching emoji used by the Emotion Hub.",
    "Mood-aware colour palette and emoji mapping.",
)

H3("Reusable Widgets")
P(
    "Five reusable widgets carry the bulk of the visual workload. MainShell is "
    "the persistent bottom-navigation scaffold that hosts the five primary "
    "tabs. GlassCard is a translucent rounded-corner card with a subtle "
    "backdrop blur, used as the visual base of dashboards, automation rows, and "
    "Spotify track tiles. SensorTile and SensorCard are sensor reading "
    "containers with threshold-driven colour coding (gas, vibration, "
    "temperature). CustomButton provides a primary action button with "
    "consistent height, padding, and disabled-state styling. This small set of "
    "building blocks allows new screens to be assembled rapidly while "
    "maintaining strict visual coherence across the application."
)


# ============================================================================
# 8. External REST API Catalogue (consolidated reference)
# ============================================================================
H2("External REST API Catalogue")
P(
    "The mobile application interacts with three external services in addition "
    "to the AWS API Gateway endpoints documented earlier in this chapter. "
    "Table 5.8 consolidates the complete off-cloud call surface so that the "
    "boundaries between the application, the home network, and external "
    "providers can be reviewed in one place."
)
CAPTION("Table 5.8: External REST endpoints invoked by the application")
TABLE(
    ["Service", "Method + URL", "Used by", "Authentication"],
    [
        ["Anthropic Messages", "POST https://api.anthropic.com/v1/messages", "AiAgentService (chat agent), SpotifyClaudeAgent (music curator)", "x-api-key header from .env"],
        ["Spotify Accounts", "POST https://accounts.spotify.com/api/token", "SpotifyAuth.exchangeCodeForToken / refresh", "Basic <base64(client_id:client_secret)>"],
        ["Spotify Web API", "GET https://api.spotify.com/v1/me/top/tracks", "SpotifyApiClient.getTopTracks", "Bearer <access_token>"],
        ["Spotify Web API", "GET https://api.spotify.com/v1/me/top/artists", "SpotifyApiClient.getTopArtists", "Bearer <access_token>"],
        ["Spotify Web API", "GET https://api.spotify.com/v1/me/player/recently-played", "SpotifyApiClient.getRecentlyPlayed", "Bearer <access_token>"],
        ["Spotify Web API", "GET https://api.spotify.com/v1/search", "SpotifyApiClient.searchTrack (Claude resolver)", "Bearer <access_token>"],
        ["Raspberry Pi", "POST https://<pi_host>:8000/predict", "EmotionApiService.predictFromJpeg", "TLS only (LAN-restricted)"],
    ],
)
P(
    "All AWS endpoints invoked by the application are documented in Tables 5.1 "
    "and 5.2 earlier in this chapter, with the addition of POST and DELETE "
    "/{home_id}/automations introduced for the automations module. Together, "
    "Tables 5.1, 5.2 and 5.8 form a complete inventory of every network call "
    "the application is capable of making."
)


# ============================================================================
# 9. Configuration, Testing, and CI
# ============================================================================
H2("Configuration, Testing, and Continuous Integration")

H3("Environment Variable Management")
P(
    "Secrets are kept out of source control through flutter_dotenv. A .env file "
    "(listed in .gitignore) holds four variables: CLAUDE_API_KEY, "
    "SPOTIFY_CLIENT_ID, SPOTIFY_CLIENT_SECRET, and an optional EMOTION_API_URL "
    "override. The file is bundled as a Flutter asset so it is read at runtime "
    "rather than at build time, allowing the same compiled binary to be "
    "redeployed with different Pi addresses across staging and production. "
    "Firebase configuration is generated by the FlutterFire CLI into "
    "firebase_options.dart, and the Cognito configuration lives in "
    "amplifyconfiguration.dart produced by amplify pull — both of these files "
    "contain non-secret identifiers and are therefore safe to commit."
)

H3("Widget and Unit Testing")
P(
    "The test/ directory contains widget tests written against flutter_test. "
    "The tests exercise the rendering paths of dependency-free widgets "
    "(SensorTile, CustomButton) and stub the network layer through Riverpod "
    "overrides so that screens under test never make real HTTP calls. The "
    "mood-keyword matcher in SpotifyMoodCatalog is covered by pure-Dart unit "
    "tests that assert the expected scoring behaviour for representative "
    "Turkish and English track names."
)

H3("Continuous Integration")
P(
    "A GitHub Actions workflow at .github/workflows/flutter-ci.yml runs on "
    "every push and pull request. The workflow installs the Flutter SDK at the "
    "version pinned in pubspec.lock, executes flutter pub get, runs flutter "
    "analyze with the lint rules defined in analysis_options.yaml, and finally "
    "executes flutter test. A green status check is required before any change "
    "can be merged into main, providing an automated guard against regression "
    "of either lints or test coverage."
)
FIG(
    "5.26",
    "GitHub Actions run page for the flutter-ci.yml workflow showing a "
    "successful pipeline with green checkmarks on the analyze and test "
    "steps, including timing information for each step.",
    "GitHub Actions CI run for the Flutter pipeline.",
)

H3("Build Configuration and Deployment")
P(
    "The Android build is configured in android/app/build.gradle with "
    "minSdkVersion 21 (sufficient for Amplify and Firebase) and "
    "targetSdkVersion 34. The iOS build uses an iOS 13.0 deployment target "
    "driven by the firebase_messaging plugin's minimum requirement. Release "
    "builds are produced through flutter build apk --release and flutter build "
    "ipa --release; signing is performed with a project-specific keystore for "
    "Android and the team's distribution certificate for iOS. The redirect URI "
    "akilliev://callback used by the Spotify OAuth flow is registered as an "
    "Android intent filter and as an iOS CFBundleURLSchemes entry so that the "
    "OS routes the post-authorisation redirect back into the app."
)


# ============================================================================
# 10. Known Limitations and Production Considerations
# ============================================================================
H2("Known Limitations and Production Considerations")
P(
    "While the implementation satisfies all functional requirements defined in "
    "Chapter 3, several constraints arise from external platform decisions and "
    "free-tier ceilings rather than from the system design itself, and are "
    "documented here for transparency."
)
P(
    "Spotify deprecations. The November 2024 removal of /v1/recommendations, "
    "/v1/audio-features, and /v1/artists/{id}/top-tracks for newly created apps "
    "eliminated the originally planned audio-feature-driven mood matching. The "
    "keyword-based catalog pipeline and the Claude-curated candidate pipeline "
    "together provide a resilient replacement that remains entirely personal, "
    "but they cannot recover the granular tempo / valence / energy targets the "
    "deprecated endpoints offered. Mood matching is therefore coarser than the "
    "original design envisaged."
)
P(
    "Self-signed Pi certificate. The development setup uses a self-signed TLS "
    "certificate on the Raspberry Pi, with a debug-only certificate-validation "
    "bypass in the mobile client. A production rollout requires replacing the "
    "self-signed certificate with a Let's Encrypt or private-CA certificate and "
    "switching EmotionApiService._buildClient() back to a plain http.Client() "
    "that performs full validation."
)
P(
    "Network resilience. AWS execute-api errors are deliberately suppressed "
    "after the first log line so that an intermittent Wi-Fi outage does not "
    "flood the developer console. The trade-off is that any sustained backend "
    "failure is silent at the application layer until the network recovers; "
    "users instead notice this through the dashboard's stale-data dimming and "
    "through the absence of incoming alerts. This is acceptable for the current "
    "prototype scope but would be replaced with a user-visible offline banner "
    "in a commercial deployment."
)
P(
    "Emulator face detection. Android emulators present a synthetic camera feed "
    "that does not contain a real face, so the Pi model returns no_face_detected "
    "whenever a scan is attempted from the emulator. Development testing of the "
    "emotion pipeline therefore requires either a physical device, the manual "
    "mood picker, or the ImageSource.gallery substitution with a sample face "
    "photograph."
)
P(
    "Cloud free-tier ceilings. AWS IoT Core, Lambda invocation count, and RDS "
    "storage are all consumed under free-tier limits. The current workload "
    "remains comfortably below these thresholds, but a multi-home rollout "
    "beyond a small pilot would require migrating to paid tiers — an exercise "
    "outside the scope of this thesis as defined in Chapter 1."
)


# ---------------------------------------------------------------------------
# Apply
# ---------------------------------------------------------------------------

def main():
    doc = Document(SRC)

    # Phase 1 — patch the existing chapter (Table 5.2 + new endpoints)
    patch_existing_chapter(doc)

    # Phase 2 — append the new chapter content before REFERENCES
    anchor = find_anchor(doc)

    fig_count = 0
    table_count = 0

    for kind, payload in SECTIONS:
        if kind == "__TABLE__":
            headers, rows = payload
            insert_table(doc, anchor, headers, rows)
            table_count += 1
        elif kind == "__FIG__":
            number, description, caption, fig_kind = payload
            insert_figure_placeholder(
                doc, anchor, number, description, caption, fig_kind
            )
            fig_count += 1
        elif kind == "__IMG__":
            number, image_path, caption = payload
            insert_image(doc, anchor, image_path, number, caption)
            fig_count += 1
        elif kind == "__CODE__":
            insert_paragraph_at(doc, anchor, payload, style="Normal", role="code")
        elif kind == "Caption":
            insert_paragraph_at(doc, anchor, payload, style="Caption", role="caption")
        else:
            insert_paragraph_at(doc, anchor, payload, style=kind, role="body")

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"  Inserted {fig_count} figures")
    print(f"  Inserted {table_count} new tables in new chapter")
    print(f"  Patched existing chapter: 2 rows added to Table 5.2,")
    print(f"  + 2 endpoint description blocks before Edge Layer & MQTT section")


if __name__ == "__main__":
    main()
